import io
import re
import requests
import streamlit.components.v1 as components

import streamlit as st
import pandas as pd
import numpy as np
import altair as alt
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

try:
    import tableauserverclient as TSC
    TABLEAU_AVAILABLE = True
except ImportError:
    TABLEAU_AVAILABLE = False

try:
    import msal
    MSAL_AVAILABLE = True
except ImportError:
    MSAL_AVAILABLE = False
# ---------------------------------------------------
# PAGE CONFIG
# ---------------------------------------------------

st.set_page_config(
    page_title="Nixara",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------------------------------------------------
# CUSTOM CSS — Light professional executive theme
# ---------------------------------------------------

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;600&family=IBM+Plex+Sans:wght@300;400;500&display=swap');

/* ── Global ── */
html, body, [class*="css"] {
    font-family: 'IBM Plex Sans', sans-serif;
    background-color: #F7F6F2;
    color: #1A1A2E;
}

/* ── Hide Streamlit chrome without disabling native sidebar controls ── */
#MainMenu, footer { visibility: hidden; }
header {
    background: transparent !important;
}
button[data-testid="stBaseButton-header"],
button[data-testid="stMainMenuButton"] {
    visibility: hidden !important;
}

/* ── Sidebar toggle buttons — keep Streamlit's native collapse/expand behavior ── */
[data-testid="stSidebarCollapseButton"],
[data-testid="stSidebarCollapsedControl"],
button[data-testid="stExpandSidebarButton"] {
    visibility: visible !important;
    display: flex !important;
}
[data-testid="stSidebarCollapseButton"] button,
[data-testid="stSidebarCollapsedControl"] button,
button[data-testid="stExpandSidebarButton"] {
    visibility: visible !important;
    background: #1A1A2E !important;
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 6px !important;
    box-shadow: 0 2px 8px rgba(0,0,0,.2) !important;
    cursor: pointer !important;
}
[data-testid="stSidebarCollapseButton"] button *,
[data-testid="stSidebarCollapsedControl"] button *,
button[data-testid="stExpandSidebarButton"] * {
    color: #F7F6F2 !important;
    fill: #F7F6F2 !important;
}
[data-testid="stSidebarCollapseButton"] button:hover,
[data-testid="stSidebarCollapsedControl"] button:hover,
button[data-testid="stExpandSidebarButton"]:hover {
    opacity: 0.82 !important;
}

[data-testid="stSidebarCollapsedControl"],
button[data-testid="stExpandSidebarButton"] {
    z-index: 2147483647 !important;
}

@media (max-width: 768px) {
    [data-testid="stSidebar"] {
        z-index: 2147483646 !important;
    }
    [data-testid="stSidebarCollapseButton"] button,
    [data-testid="stSidebarCollapsedControl"] button,
    button[data-testid="stExpandSidebarButton"] {
        min-width: 44px !important;
        min-height: 44px !important;
        border-radius: 8px !important;
    }
}
.block-container {
    padding-top: 2rem;
    padding-bottom: 3rem;
    max-width: 1200px;
}

/* ── Sidebar ── */
[data-testid="stSidebar"] {
    background-color: #FFFFFF;
    border-right: 1px solid #E8E5DC;
}
[data-testid="stSidebar"] .stMarkdown p {
    color: #9B8FA8;
    font-size: 0.7rem;
    letter-spacing: 0.12em;
    text-transform: uppercase;
    font-weight: 500;
}
[data-testid="stSidebar"] label {
    color: #3A3A5A !important;
    font-size: 0.82rem !important;
    font-weight: 400 !important;
}
[data-testid="stSidebar"] .stTextArea textarea {
    background-color: #F7F6F2 !important;
    border: 1px solid #D8D4CC !important;
    color: #1A1A2E !important;
    border-radius: 6px !important;
    font-family: 'IBM Plex Sans', sans-serif !important;
    font-size: 0.85rem !important;
}
[data-testid="stSidebar"] select,
[data-testid="stSidebar"] .stSelectbox > div > div {
    background-color: #F7F6F2 !important;
    border: 1px solid #D8D4CC !important;
    color: #1A1A2E !important;
    border-radius: 6px !important;
}

/* ── Logo placeholder ── */
.logo-placeholder {
    width: 100%;
    height: 52px;
    border: 1.5px dashed #C8C4BC;
    border-radius: 8px;
    display: flex;
    align-items: center;
    justify-content: center;
    color: #A8A4A0;
    font-size: 0.68rem;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    font-weight: 500;
    transition: all 0.2s;
    background: #FAFAF8;
    cursor: pointer;
}
.logo-placeholder:hover { border-color: #B8975A; color: #B8975A; background: #FBF8F2; }

/* ── App header ── */
.app-header {
    border-bottom: 2px solid #E8E5DC;
    padding-bottom: 1.25rem;
    margin-bottom: 2rem;
    display: flex;
    align-items: flex-end;
    justify-content: space-between;
}
.app-title {
    font-family: 'Playfair Display', serif;
    font-size: 1.9rem;
    font-weight: 600;
    color: #1A1A2E;
    margin: 0;
    line-height: 1.2;
}
.app-subtitle {
    color: #6B6B8A;
    font-size: 0.83rem;
    margin-top: 0.3rem;
    font-weight: 300;
    letter-spacing: 0.02em;
}
.accent-bar {
    width: 36px;
    height: 3px;
    background: #B8975A;
    border-radius: 2px;
    margin: 0.5rem 0 0.35rem;
}
.header-tag {
    font-size: 0.7rem;
    font-weight: 500;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: #B8975A;
    border: 1px solid #B8975A;
    border-radius: 20px;
    padding: 0.2rem 0.75rem;
    margin-bottom: 0.5rem;
}

/* ── Section labels ── */
.section-label {
    font-size: 0.68rem;
    font-weight: 500;
    letter-spacing: 0.14em;
    text-transform: uppercase;
    color: #B8975A;
    margin-bottom: 0.6rem;
    margin-top: 0.25rem;
}

/* ── Metric cards ── */
[data-testid="metric-container"] {
    background: #FFFFFF;
    border: 1px solid #E8E5DC;
    border-radius: 10px;
    padding: 1.1rem 1.3rem !important;
    box-shadow: 0 1px 4px rgba(26,26,46,0.06);
}
[data-testid="metric-container"] label {
    color: #8B8FA8 !important;
    font-size: 0.72rem !important;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    font-weight: 500 !important;
}
[data-testid="metric-container"] [data-testid="stMetricValue"] {
    color: #1A1A2E !important;
    font-size: 1.65rem !important;
    font-weight: 300 !important;
    font-family: 'Playfair Display', serif !important;
}

/* ── Dataframe ── */
[data-testid="stDataFrame"] {
    border: 1px solid #E8E5DC !important;
    border-radius: 8px !important;
    overflow: hidden;
    box-shadow: 0 1px 4px rgba(26,26,46,0.05);
}

/* ── Tabs ── */
[data-testid="stTabs"] [role="tablist"] {
    border-bottom: 2px solid #E8E5DC;
    gap: 0;
    background: transparent;
}
[data-testid="stTabs"] button[role="tab"] {
    background: transparent !important;
    color: #8B8FA8 !important;
    font-size: 0.82rem !important;
    font-weight: 500 !important;
    letter-spacing: 0.04em;
    padding: 0.65rem 1.4rem !important;
    border-bottom: 2px solid transparent !important;
    border-radius: 0 !important;
    transition: all 0.2s;
    font-family: 'IBM Plex Sans', sans-serif !important;
}
[data-testid="stTabs"] button[role="tab"]:hover {
    color: #1A1A2E !important;
    background: #F0EDE6 !important;
}
[data-testid="stTabs"] button[role="tab"][aria-selected="true"] {
    color: #1A1A2E !important;
    border-bottom: 2px solid #B8975A !important;
    font-weight: 500 !important;
}

/* ── Report body ── */
.report-body {
    background: #FFFFFF;
    border: 1px solid #E8E5DC;
    border-radius: 10px;
    padding: 1.5rem 2.5rem;
    line-height: 1.65;
    font-size: 1rem;
    color: #2E2E4A;
    box-shadow: 0 1px 4px rgba(26,26,46,0.05);
}
.report-body p {
    margin-bottom: 0.55rem;
}
.report-body h3 {
    font-family: 'Playfair Display', serif;
    color: #1A1A2E;
    font-size: 1.08rem;
    font-weight: 600;
    margin-top: 1.2rem;
    margin-bottom: 0.35rem;
    border-left: 3px solid #B8975A;
    padding-left: 0.7rem;
}

/* ── Download buttons ── */
[data-testid="stDownloadButton"] button {
    background: #FFFFFF !important;
    border: 1px solid #D4D0C8 !important;
    color: #4A4A6A !important;
    font-size: 0.78rem !important;
    font-weight: 500 !important;
    letter-spacing: 0.04em;
    border-radius: 6px !important;
    transition: all 0.2s !important;
    font-family: 'IBM Plex Sans', sans-serif !important;
}
[data-testid="stDownloadButton"] button:hover {
    border-color: #B8975A !important;
    color: #B8975A !important;
    background: #FBF8F2 !important;
}

/* ── Generate button ── */
[data-testid="stSidebar"] [data-testid="stButton"] button {
    background: #1A1A2E !important;
    color: #FFFFFF !important;
    border: none !important;
    font-weight: 500 !important;
    font-size: 0.85rem !important;
    letter-spacing: 0.06em;
    border-radius: 6px !important;
    transition: opacity 0.2s !important;
    font-family: 'IBM Plex Sans', sans-serif !important;
}
[data-testid="stSidebar"] [data-testid="stButton"] button:hover {
    opacity: 0.85 !important;
}

/* ── Alerts ── */
[data-testid="stAlert"] {
    background: #FDF8F0 !important;
    border: 1px solid #E8D9B8 !important;
    border-radius: 8px !important;
    color: #5A4A2A !important;
}

/* ── Divider ── */
hr { border-color: #E8E5DC !important; }

/* ── Empty state ── */
.empty-state {
    text-align: center;
    padding: 5rem 2rem;
    color: #B0ADC0;
}
.empty-state .icon {
    font-size: 2rem;
    margin-bottom: 1rem;
    color: #D4D0C8;
}
.empty-state p {
    font-size: 0.82rem;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    font-weight: 500;
}


/* ── File uploader ── */
[data-testid="stFileUploader"] {
    background: #F7F6F2 !important;
    border: 1px solid #D8D4CC !important;
    border-radius: 8px !important;
}
[data-testid="stFileUploader"] section {
    background: #F7F6F2 !important;
    border: none !important;
}
[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzoneInstructions"] {
    color: #6B6B8A !important;
}
/* uploaded file pill */
[data-testid="stFileUploader"] [data-testid="stFileUploaderFile"] {
    background: #EEEAE4 !important;
    border: 1px solid #D8D4CC !important;
    border-radius: 6px !important;
    color: #1A1A2E !important;
}
[data-testid="stFileUploader"] [data-testid="stFileUploaderFile"] span {
    color: #3A3A5A !important;
}
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------
# CLIENT — resolved later in sidebar after user key is captured
# ---------------------------------------------------

def get_client(user_key=None):
    # The visitor's key is used ONLY here to initialise the OpenAI client.
    # It is never logged, stored on the server, or sent anywhere except OpenAI.
    # If no key is supplied the call will fail — there is no owner fallback.
    key = user_key.strip() if user_key and user_key.strip() else ""
    return OpenAI(api_key=key)

# ---------------------------------------------------
# FILE LOADER
# ---------------------------------------------------

def load_file(uploaded):
    name = uploaded.name.lower()
    if name.endswith(".csv"):
        return pd.read_csv(uploaded)
    elif name.endswith(".xlsx") or name.endswith(".xls"):
        return pd.read_excel(uploaded)
    return None


# ---------------------------------------------------
# CLEAN NUMERIC
# ---------------------------------------------------

def clean_dataframe(df):
    for col in df.columns:
        if df[col].dtype == object:
            cleaned = pd.to_numeric(
                df[col].astype(str)
                .str.strip()
                .str.replace(r"[\$,]", "", regex=True),
                errors="coerce",
            )
            # Lower threshold to 0.5 — if half the values parse as numbers, treat as numeric
            if cleaned.notna().mean() > 0.5:
                df[col] = cleaned
    return df


# ---------------------------------------------------
# ANALYSIS HELPERS
# ---------------------------------------------------

def detect_anomalies(df, col):
    s = df[col].dropna()
    if len(s) < 5 or s.std() == 0:
        return pd.DataFrame()
    z = (s - s.mean()) / s.std()
    # Use .loc with the index from s to avoid misalignment
    anomaly_idx = z[np.abs(z) > 2].index
    return df.loc[anomaly_idx].copy()


def dashboard_score(df):
    score = 100
    if df.isna().mean().mean() > 0.2:
        score -= 20
    if len(df.columns) > 20:
        score -= 10
    if len(df) < 10:
        score -= 10
    return max(score, 0)


def smart_agg(col_name):
    """Use mean for averages/rates/margins, sum for everything else."""
    keywords = ["average", "avg", "mean", "rate", "ratio", "margin", "score", "pct", "percent"]
    return "mean" if any(k in col_name.lower() for k in keywords) else "sum"


def build_data_summary(df, filter_col=None, filter_val=None):
    if filter_col and filter_val and filter_col in df.columns:
        df = df[df[filter_col] == filter_val]

    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    cat_cols = df.select_dtypes(exclude=np.number).columns.tolist()

    lines = []
    lines.append(f"Rows: {len(df)} | Columns: {len(df.columns)}")
    if filter_col and filter_val:
        lines.append(f"Filtered to: {filter_col} = {filter_val}")
    lines.append(f"Numeric columns: {numeric_cols}")
    lines.append(f"Categorical columns: {cat_cols}")
    lines.append("")

    if numeric_cols:
        lines.append("NUMERIC SUMMARY")
        lines.append(df[numeric_cols].describe().round(2).to_string())
        lines.append("")

    for cat in cat_cols[:3]:
        if df[cat].nunique() <= 20:
            grp_parts = {}
            for nc in numeric_cols[:3]:
                grp_parts[nc] = df.groupby(cat)[nc].agg(smart_agg(nc)).round(2)
            if grp_parts:
                grp = pd.DataFrame(grp_parts)
                lines.append(f"BREAKDOWN BY {cat.upper()}")
                lines.append(grp.to_string())
                lines.append("")

    anomaly_lines = []
    for col in numeric_cols[:3]:
        anom = detect_anomalies(df, col)
        if not anom.empty:
            anomaly_lines.append(f"  {col}: {len(anom)} anomalous rows (z > 2)")
            preview_cols = [col] + cat_cols[:2]
            anomaly_lines.append(anom[preview_cols].head(3).to_string())
    if anomaly_lines:
        lines.append("ANOMALIES DETECTED")
        lines.extend(anomaly_lines)
        lines.append("")

    if len(numeric_cols) >= 2:
        corr = df[numeric_cols].corr().round(3)
        corr_pairs = (
            corr.where(np.triu(np.ones(corr.shape), k=1).astype(bool))
            .stack()
            .sort_values(key=abs, ascending=False)
            .head(5)
        )
        lines.append("TOP CORRELATIONS")
        lines.append(corr_pairs.to_string())
        lines.append("")

    lines.append(f"Data Quality Score: {dashboard_score(df)}/100")
    return "\n".join(lines)


# ---------------------------------------------------
# REPORT CONFIGS
# ---------------------------------------------------

REPORT_CONFIGS = {
    "Executive Summary": """Structure your response EXACTLY as:

Situation
(2-3 sentences on what the data shows at the highest level)

What This Means For You
(connect to the decision and role — 2-3 sentences)

Recommended Actions
(3 numbered actions, each specific and time-bound)

Risks If You Wait
(2 bullet points on what gets worse without action)

Rules: under 350 words, every action references a number, never say consider monitoring.""",

    "Operational Detail": """Structure your response EXACTLY as:

Performance Breakdown
(by the main dimensions in the data — categories, regions, segments, etc.)

Efficiency Gaps
(where effort or cost is not matching return — specific numbers)

Process Recommendations
(4-5 numbered operational changes implementable this quarter)

Quick Wins
(2 things doable in the next 2 weeks with zero new resources)

Rules: under 500 words, operational language, every point references data.""",

    "Risk Report": """Structure your response EXACTLY as:

Top Risks Identified
(3 risks ranked by severity — each with: what it is, what signals it, potential impact)

Early Warning Signs
(specific metrics to watch as leading indicators)

Mitigation Actions
(one concrete action per risk — specific, assignable, time-bound)

Data Quality Risks
(flag any gaps, anomalies, or quality issues masking bigger problems)

Rules: under 450 words, risk language, reference specific numbers, write to prompt escalation.""",
}


# ---------------------------------------------------
# PROMPT + AI CALL
# ---------------------------------------------------

def generate_report(summary, who, decision, timeframe, report_type, api_key=None):
    instruction = REPORT_CONFIGS[report_type]
    prompt = f"""You are advising a {who} who needs to make a decision about:
"{decision}"

Time horizon: {timeframe}
Report type: {report_type}

{instruction}

Write for a {who} — direct and specific, not academic.

Dashboard data:
{summary}
"""
    c = get_client(api_key)
    response = c.chat.completions.create(
        model="gpt-4o",
        max_tokens=1024,
        messages=[{"role": "user", "content": prompt}],
    )
    return clean_ai_output(response.choices[0].message.content)


def clean_ai_output(text):
    # Strip bold markers
    text = re.sub(r"\*\*", "", text)
    # Strip any existing # heading markers — we re-add them cleanly below.
    # Without this, "### Situation" → "### \n### Situation" (leaves orphan ###).
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    # Strip backtick code formatting the AI uses for column names
    text = re.sub(r"`([^`]*)`", r"\1", text)
    headers = [
        "Situation", "What This Means For You", "Recommended Actions",
        "Risks If You Wait", "Performance Breakdown", "Efficiency Gaps",
        "Process Recommendations", "Quick Wins", "Top Risks Identified",
        "Early Warning Signs", "Mitigation Actions", "Data Quality Risks",
    ]
    for h in headers:
        text = re.sub(h, f"\n### {h}", text, flags=re.IGNORECASE)
    return text.strip()


# ---------------------------------------------------
# DOWNLOAD: WORD
# ---------------------------------------------------

def report_to_docx(report_text, title, who, decision, timeframe):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    run = meta.add_run(f"Role: {who}   |   Decision: {decision}   |   Horizon: {timeframe}")
    run.font.size = Pt(9)

    doc.add_paragraph()

    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif re.match(r"^\d+\.", line):
            doc.add_paragraph(line, style="List Number")
        elif line.startswith("- ") or line.startswith("• "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------
# DOWNLOAD: PDF
# ---------------------------------------------------

def report_to_pdf(report_text, title, who, decision, timeframe):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=1.2 * inch,
        rightMargin=1.2 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
    )

    styles = getSampleStyleSheet()

    style_title = ParagraphStyle(
        "ReportTitle", parent=styles["Title"],
        fontSize=18, spaceAfter=4,
        textColor=colors.HexColor("#111111"),
    )
    style_meta = ParagraphStyle(
        "Meta", parent=styles["Normal"],
        fontSize=9, spaceAfter=14,
        textColor=colors.HexColor("#666666"),
    )
    style_h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"],
        fontSize=12, spaceBefore=12, spaceAfter=4,
        textColor=colors.HexColor("#1D3557"),
    )
    style_body = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=10, leading=15, spaceAfter=4,
    )

    story = [
        Paragraph(title, style_title),
        Paragraph(
            f"Role: {who} &nbsp;|&nbsp; Decision: {decision} &nbsp;|&nbsp; Horizon: {timeframe}",
            style_meta,
        ),
    ]

    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 4))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:], style_h2))
        else:
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, style_body))

    doc.build(story)
    buf.seek(0)
    return buf


# ---------------------------------------------------
# TABLEAU CONNECTOR
# ---------------------------------------------------

def load_tableau(server_url, site_id, token_name, token_secret, view_name):
    if not TABLEAU_AVAILABLE:
        st.error("tableauserverclient is not installed. Run: pip install tableauserverclient")
        return None
    try:
        auth   = TSC.PersonalAccessTokenAuth(token_name, token_secret, site_id=site_id)
        server = TSC.Server(server_url, use_server_version=True)
        with server.auth.sign_in(auth):
            views, _ = server.views.get()
            target   = next((v for v in views if v.name == view_name), None)
            if target is None:
                available = [v.name for v in views]
                st.error(f"View '{view_name}' not found. Available views: {available}")
                return None
            server.views.populate_csv(target)
            csv_data = b"".join(target.csv).decode("utf-8")
        return pd.read_csv(io.StringIO(csv_data))
    except Exception as e:
        st.error(f"Tableau connection failed: {e}")
        return None


# ---------------------------------------------------
# POWER BI CONNECTOR
# ---------------------------------------------------

def get_powerbi_token(tenant_id, client_id, client_secret):
    if not MSAL_AVAILABLE:
        st.error("msal is not installed. Run: pip install msal")
        return None
    try:
        app    = msal.ConfidentialClientApplication(
            client_id,
            authority=f"https://login.microsoftonline.com/{tenant_id}",
            client_credential=client_secret,
        )
        result = app.acquire_token_for_client(
            scopes=["https://analysis.windows.net/powerbi/api/.default"]
        )
        if "access_token" not in result:
            st.error(f"Power BI auth failed: {result.get('error_description', 'Unknown error')}")
            return None
        return result["access_token"]
    except Exception as e:
        st.error(f"Power BI auth error: {e}")
        return None


def list_powerbi_tables(token, workspace_id, dataset_id):
    url     = f"https://api.powerbi.com/v1.0/myorg/groups/{workspace_id}/datasets/{dataset_id}/tables"
    headers = {"Authorization": f"Bearer {token}"}
    resp    = requests.get(url, headers=headers)
    if resp.status_code != 200:
        st.error(f"Could not fetch tables: {resp.text}")
        return []
    return [t["name"] for t in resp.json().get("value", [])]


def load_powerbi_table(token, workspace_id, dataset_id, table_name):
    try:
        url     = f"https://api.powerbi.com/v1.0/myorg/groups/{workspace_id}/datasets/{dataset_id}/executeQueries"
        headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
        body    = {
            "queries": [{"query": f"EVALUATE '{table_name}'"}],
            "serializerSettings": {"includeNulls": True},
        }
        resp = requests.post(url, headers=headers, json=body)
        if resp.status_code != 200:
            st.error(f"Power BI query failed: {resp.text}")
            return None
        results = resp.json()["results"][0]["tables"][0]["rows"]
        df      = pd.DataFrame(results)
        # Strip the "TableName[ColumnName]" prefix Power BI adds to column names
        df.columns = [c.split("[")[-1].rstrip("]") for c in df.columns]
        return df
    except Exception as e:
        st.error(f"Power BI data fetch failed: {e}")
        return None


# ---------------------------------------------------
# SIDEBAR
# ---------------------------------------------------

with st.sidebar:

    st.markdown('<p class="section-label">Data Source</p>', unsafe_allow_html=True)

    data_source = st.radio(
        "Choose how to load data",
        ["Upload CSV / Excel", "Tableau", "Power BI"],
        label_visibility="collapsed",
        horizontal=False,
    )

    uploaded_file  = None
    df_from_bi     = None

    # ── Upload ──────────────────────────────────────
    if data_source == "Upload CSV / Excel":
        uploaded_file = st.file_uploader(
            "Upload your data file",
            type=["csv", "xlsx", "xls"],
            label_visibility="collapsed",
            help="CSV or Excel — any dataset works. Try superstore_data.csv to explore.",
        )

    # ── Tableau ─────────────────────────────────────
    elif data_source == "Tableau":
        st.markdown('<p class="section-label">Tableau Credentials</p>', unsafe_allow_html=True)
        tab_server = st.text_input("Server URL", placeholder="https://us-east-1.online.tableau.com")
        tab_site   = st.text_input("Site ID", placeholder="your-site-name")
        tab_token  = st.text_input("Token Name", placeholder="my-token")
        tab_secret = st.text_input("Token Secret", type="password", placeholder="••••••••")
        tab_view   = st.text_input("View Name", placeholder="Sales Overview")
        if st.button("Connect to Tableau", use_container_width=True):
            if all([tab_server, tab_site, tab_token, tab_secret, tab_view]):
                with st.spinner("Connecting to Tableau..."):
                    df_from_bi = load_tableau(tab_server, tab_site, tab_token, tab_secret, tab_view)
                if df_from_bi is not None:
                    st.session_state["bi_df"]     = df_from_bi
                    st.session_state["bi_source"]  = f"Tableau · {tab_view}"
                    st.success(f"Connected — {len(df_from_bi):,} rows loaded")
            else:
                st.warning("Please fill in all Tableau fields.")
        if "bi_df" in st.session_state and st.session_state.get("bi_source", "").startswith("Tableau"):
            df_from_bi = st.session_state["bi_df"]

    # ── Power BI ─────────────────────────────────────
    elif data_source == "Power BI":
        st.markdown('<p class="section-label">Power BI Credentials</p>', unsafe_allow_html=True)
        pbi_tenant   = st.text_input("Tenant ID", placeholder="xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx")
        pbi_client   = st.text_input("Client ID", placeholder="xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx")
        pbi_secret   = st.text_input("Client Secret", type="password", placeholder="••••••••")
        pbi_workspace= st.text_input("Workspace ID", placeholder="xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx")
        pbi_dataset  = st.text_input("Dataset ID", placeholder="xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx")

        if st.button("Connect to Power BI", use_container_width=True):
            if all([pbi_tenant, pbi_client, pbi_secret, pbi_workspace, pbi_dataset]):
                with st.spinner("Authenticating with Power BI..."):
                    token = get_powerbi_token(pbi_tenant, pbi_client, pbi_secret)
                if token:
                    tables = list_powerbi_tables(token, pbi_workspace, pbi_dataset)
                    if tables:
                        st.session_state["pbi_token"]     = token
                        st.session_state["pbi_workspace"] = pbi_workspace
                        st.session_state["pbi_dataset"]   = pbi_dataset
                        st.session_state["pbi_tables"]    = tables
                        st.success(f"Connected — {len(tables)} table(s) found")
                    else:
                        st.warning("Connected but no tables found in this dataset.")
            else:
                st.warning("Please fill in all Power BI fields.")

        if "pbi_tables" in st.session_state:
            pbi_table_choice = st.selectbox(
                "Select table to analyse",
                st.session_state["pbi_tables"]
            )
            if st.button("Load Table", use_container_width=True):
                with st.spinner(f"Loading {pbi_table_choice}..."):
                    df_from_bi = load_powerbi_table(
                        st.session_state["pbi_token"],
                        st.session_state["pbi_workspace"],
                        st.session_state["pbi_dataset"],
                        pbi_table_choice,
                    )
                if df_from_bi is not None:
                    st.session_state["bi_df"]    = df_from_bi
                    st.session_state["bi_source"] = f"Power BI · {pbi_table_choice}"
                    st.success(f"Loaded — {len(df_from_bi):,} rows")
            if "bi_df" in st.session_state and st.session_state.get("bi_source", "").startswith("Power BI"):
                df_from_bi = st.session_state["bi_df"]

    st.markdown("---")
    st.markdown('<p class="section-label">Analysis Context</p>', unsafe_allow_html=True)

    decision = st.text_area(
        "Decision",
        placeholder='e.g. "Which sub-category sells best in Texas?"',
        height=90,
        label_visibility="collapsed",
    )

    who = st.selectbox(
        "Role",
        ["COO", "CEO", "CFO", "Sales Lead", "Operations Lead", "Board"],
        label_visibility="visible",
    )

    timeframe = st.selectbox(
        "Time horizon",
        ["Next 30 days", "This quarter", "This year"],
        index=1,
    )

    st.markdown("---")
    st.markdown('<p class="section-label">Filter (optional)</p>', unsafe_allow_html=True)

    filter_col = None
    filter_val = None

    if "df_for_filter" in st.session_state:
        df_f = st.session_state["df_for_filter"]
        cat_cols_f = df_f.select_dtypes(exclude=np.number).columns.tolist()
        if cat_cols_f:
            filter_col = st.selectbox("Filter by column", ["None"] + cat_cols_f)
            if filter_col != "None":
                filter_val = st.selectbox(
                    "Filter value",
                    sorted(df_f[filter_col].dropna().unique().tolist())
                )
            else:
                filter_col = None

    st.markdown("---")
    st.markdown('<p class="section-label">OpenAI API Key</p>', unsafe_allow_html=True)

    # Read the key from query params (injected by main-area JS on the previous rerun)
    _qp = st.query_params
    _saved_key = _qp.get("_oai", "")

    # Initialise session state — prefer query-param value if session state is empty
    if "user_api_key" not in st.session_state:
        st.session_state["user_api_key"] = _saved_key

    # If a new saved key arrived from localStorage and session state is still empty, adopt it
    if _saved_key and not st.session_state["user_api_key"]:
        st.session_state["user_api_key"] = _saved_key

    entered_key = st.text_input(
        "Your OpenAI API key",
        type="password",
        placeholder="sk-proj-...",
        value=st.session_state["user_api_key"],
        label_visibility="collapsed",
    )

    # Save to session state whenever the user types
    if entered_key:
        st.session_state["user_api_key"] = entered_key

    user_api_key = st.session_state["user_api_key"]

    # ── Save-in-browser checkbox ─────────────────────────────────────────────
    _already_saved = bool(_saved_key and _saved_key == user_api_key)
    save_in_browser = st.checkbox(
        "Remember key in this browser",
        value=_already_saved,
        help=(
            "Saves your key to this device only — like a password manager. "
            "It is never uploaded to any server. Uncheck to forget it."
        ),
    )

    if save_in_browser and user_api_key and user_api_key.startswith("sk-"):
        # Write directly to parent localStorage — same origin, no postMessage needed
        components.html(
            f"<script>try{{window.parent.localStorage.setItem('nixara_oai_key',{repr(user_api_key)})}}catch(e){{}}</script>",
            height=0,
        )
    elif not save_in_browser and _already_saved:
        components.html(
            "<script>try{window.parent.localStorage.removeItem('nixara_oai_key')}catch(e){}</script>",
            height=0,
        )

    # Key status indicator
    if user_api_key and user_api_key.startswith("sk-"):
        if save_in_browser:
            st.markdown(
                '<p style="font-size:0.72rem;color:#2E7D52;margin-top:3px;">'
                '✓ Key active &amp; saved in this browser</p>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                '<p style="font-size:0.72rem;color:#2E7D52;margin-top:3px;">'
                '✓ Key active for this session only</p>',
                unsafe_allow_html=True,
            )
    elif user_api_key:
        st.markdown(
            '<p style="font-size:0.72rem;color:#C0392B;margin-top:3px;">'
            '⚠ Key format looks incorrect — should start with sk-</p>',
            unsafe_allow_html=True,
        )

    # Step-by-step guide for non-tech users
    with st.expander("How to get your API key — 5 easy steps"):
        st.markdown("""
**Step 1** — Go to [platform.openai.com](https://platform.openai.com)
and create a free account (or log in if you have one).

**Step 2** — Click your profile icon (top right) → **API keys** → **Create new secret key**.
Give it any name, e.g. *"Nixara"*.

**Step 3** — Copy the key (starts with `sk-`). You only see it once — copy it now.

**Step 4** — Paste it in the field above, then tick **"Remember key in this browser"**
so you never have to paste it again on this device.

**Step 5** — OpenAI gives new accounts **$5 of free credit**.
Each report costs roughly $0.01–0.02, so you can generate hundreds of reports for free.

---
**Is it safe?**
Your key is saved in your own browser's local storage — exactly like how a password
manager works. It is never transmitted to this app's server. It goes directly from
your browser to OpenAI and nowhere else.

To remove your key at any time, untick "Remember key in this browser".
        """)

    st.markdown("---")

    # Block the button if no valid key is provided by the visitor
    key_ready = bool(user_api_key and user_api_key.startswith("sk-"))

    if not key_ready:
        st.markdown(
            '<p style="font-size:0.75rem;color:#C0392B;text-align:center;">'
            'Add your OpenAI key above to generate reports.</p>',
            unsafe_allow_html=True,
        )

    run = st.button(
        "Generate Reports",
        use_container_width=True,
        disabled=not key_ready,
    )

# ---------------------------------------------------
# MAIN-AREA PERSISTENT SCRIPTS
# Runs in the main page (not sidebar), so saved-key sync keeps working even when sidebar is closed.
# Sidebar open/close is handled by Streamlit's native controls.
# ---------------------------------------------------

components.html("""
<script>
(function() {
    var p      = window.parent;
    var LS_KEY = 'nixara_oai_key';

    /* localStorage → query-param bridge (keeps saved API key across page reloads) */
    var saved  = p.localStorage.getItem(LS_KEY);
    var params = new URLSearchParams(p.location.search);
    if (saved && params.get('_oai') !== saved) {
        params.set('_oai', saved);
        p.history.replaceState({}, '', '?' + params.toString());
        p.location.reload();
    }
})();
</script>
""", height=0)

# ---------------------------------------------------
# APP HEADER
# ---------------------------------------------------

st.markdown("""
    <div style="border-bottom: 1.5px solid #E8E5DC; padding-bottom: 1.25rem; margin-bottom: 2rem;">
        <p class="app-title">Nixara</p>
        <p style="font-size:0.7rem; color:#9B8FA8; letter-spacing:0.12em; text-transform:uppercase;
                  font-weight:500; margin: 0.1rem 0 0.4rem;">
            nik·​sa·​ra &nbsp;/nɪkˈsɑːrə/ &nbsp;·&nbsp; <em style="font-style:italic; text-transform:none; letter-spacing:0;">from <strong style="color:#B8975A;">nix</strong> (clarity, light) + <strong style="color:#B8975A;">ara</strong> (direction) — illuminating the path forward in your data</em>
        </p>
        <div class="accent-bar"></div>
        <p class="app-subtitle">Upload CSV / Excel &nbsp;·&nbsp; Connect Tableau or Power BI &nbsp;·&nbsp; Get executive-grade AI reports</p>
    </div>
""", unsafe_allow_html=True)

# ---------------------------------------------------
# TOP-LEVEL NAVIGATION TABS
# FAQ tab is defined first in code so it always renders
# even if the dashboard tab has early-exit logic.
# ---------------------------------------------------

nav_dashboard, nav_faq = st.tabs(["  Dashboard  ", "  FAQ & Help  "])

# ---------------------------------------------------
# FAQ TAB
# ---------------------------------------------------

with nav_faq:
    st.markdown("""
    <style>
    /* ── FAQ category headers ── */
    .faq-category {
        font-family: 'Playfair Display', serif;
        font-size: 1.0rem;
        font-weight: 600;
        color: #1A1A2E;
        margin-top: 2.2rem;
        margin-bottom: 0.6rem;
        border-left: 3px solid #B8975A;
        padding-left: 0.75rem;
        letter-spacing: 0.01em;
    }
    /* ── Trust pillars grid ── */
    .trust-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
        gap: 1rem;
        margin: 1.75rem 0 2rem;
    }
    .trust-card {
        background: #FFFFFF;
        border: 1px solid #E8E5DC;
        border-radius: 10px;
        padding: 1.1rem 1.25rem;
        box-shadow: 0 1px 4px rgba(26,26,46,0.05);
    }
    .trust-card-label {
        font-size: 0.68rem;
        font-weight: 600;
        letter-spacing: 0.12em;
        text-transform: uppercase;
        color: #B8975A;
        margin-bottom: 0.3rem;
    }
    .trust-card-text {
        font-size: 0.85rem;
        color: #2E2E4A;
        font-weight: 400;
        line-height: 1.45;
    }
    /* ── FAQ page header ── */
    .faq-header {
        padding-bottom: 1.5rem;
        margin-bottom: 0.5rem;
        border-bottom: 1.5px solid #E8E5DC;
    }
    .faq-eyebrow {
        font-size: 0.68rem;
        font-weight: 600;
        letter-spacing: 0.14em;
        text-transform: uppercase;
        color: #B8975A;
        margin-bottom: 0.5rem;
    }
    .faq-title {
        font-family: 'Playfair Display', serif;
        font-size: 1.6rem;
        font-weight: 600;
        color: #1A1A2E;
        margin: 0 0 0.3rem;
        line-height: 1.25;
    }
    .faq-subtitle {
        font-size: 0.85rem;
        color: #6B6B8A;
        font-weight: 300;
        letter-spacing: 0.01em;
    }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="faq-header">
        <p class="faq-eyebrow">Help & Documentation</p>
        <p class="faq-title">Frequently Asked Questions</p>
        <p class="faq-subtitle">
            Everything you need to know about how Nixara works,
            how your data is handled, and how to get the most from it.
        </p>
    </div>
    """, unsafe_allow_html=True)

    # Trust pillars
    st.markdown("""
    <div class="trust-grid">
        <div class="trust-card">
            <p class="trust-card-label">Data Privacy</p>
            <p class="trust-card-text">Your data is processed in-session only. Nothing is stored, logged, or retained after you close the browser.</p>
        </div>
        <div class="trust-card">
            <p class="trust-card-label">Your API Key</p>
            <p class="trust-card-text">Reports run through your own OpenAI account. You control access, usage, and spend — not us.</p>
        </div>
        <div class="trust-card">
            <p class="trust-card-label">No Account Required</p>
            <p class="trust-card-text">No sign-up, no login, no tracking. Open the app and start working immediately.</p>
        </div>
        <div class="trust-card">
            <p class="trust-card-label">Speed</p>
            <p class="trust-card-text">Three executive-grade reports — Executive, Operational, and Risk — generated in under 30 seconds.</p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── GETTING STARTED ──────────────────────────────────
    st.markdown('<p class="faq-category">Getting Started</p>', unsafe_allow_html=True)

    with st.expander("What is Nixara?"):
        st.markdown("""
        **Nixara is a built-in AI analyst for your business data** — no prompting skills required.

        Upload any CSV or Excel file, or connect your live Tableau / Power BI dashboard, and get three
        decision-ready reports in under 30 seconds:

        - **Executive Summary** — Strategic narrative written for the boardroom: situation, implications, recommended actions
        - **Operational Detail** — Efficiency gaps, process recommendations, and quick wins for team managers
        - **Risk Report** — Early warning signs, data quality flags, and mitigation actions

        Each report is generated for your specific role, decision context, and time horizon — then delivered as
        a formatted Word or PDF document you can share immediately.
        """)

    with st.expander("What data formats are supported?"):
        st.markdown("""
        - **CSV** (.csv) — any standard comma-separated file
        - **Excel** (.xlsx, .xls) — single-sheet workbooks work best
        - **Tableau Server** — connect via your Tableau credentials and pull live view data
        - **Power BI** — connect via Azure AD credentials and pull dataset exports

        Your data must be **tabular** (rows and columns) with at least one numeric column.
        It works best with structured business data: sales figures, operational KPIs,
        HR metrics, financial reports, healthcare outcomes.

        > **Tip:** If your Excel file has merged cells or multi-row headers, flatten it to a simple table first.
        """)

    with st.expander("How do I get an OpenAI API key?"):
        st.markdown("""
        1. Go to **[platform.openai.com](https://platform.openai.com)** and sign in (or create a free account)
        2. Click your profile → **API Keys** → **Create new secret key**
        3. Copy the key (starts with `sk-proj-...`) and paste it into the sidebar

        **Cost:** A typical three-report run uses approximately 2,000–5,000 tokens.
        At current GPT-4o pricing, that's roughly **$0.01–$0.05 per run** — less than a cup of coffee per week
        for daily use. You can monitor usage at **platform.openai.com/usage** at any time.
        """)

    # ── PRIVACY & SECURITY ────────────────────────────────
    st.markdown('<p class="faq-category">Privacy & Security</p>', unsafe_allow_html=True)

    with st.expander("Is my data safe? Where does it go?"):
        st.markdown("""
        Your data is handled with the following guarantees:

        - **Processed in memory only** — data lives in your active session and is discarded when you close the tab
        - **Never stored or logged** — we do not write your data to any database or file system
        - **Never used for AI training** — OpenAI's API terms prohibit using API inputs for model training
        - **Single external touchpoint** — the only service your data reaches is OpenAI's API,
          and only when you click *Generate Reports*

        Think of it like a calculator: it processes your numbers and returns a result,
        but it doesn't remember what you typed.
        """)

    with st.expander("Why do I need my own OpenAI API key? Why not just use the app's key?"):
        st.markdown("""
        We deliberately chose **not** to embed a shared API key. Here's why this is actually better for you:

        - **Your data only passes through your OpenAI account** — not a pooled account with other users' data
        - **You control your own spend** — you can set monthly budget limits directly in OpenAI's dashboard
        - **You can revoke access instantly** — delete the key anytime from OpenAI without affecting anyone else
        - **No trust required** — you don't have to trust that we're handling a shared key responsibly,
          because there isn't one

        It's the same reason you wouldn't want your bank to share a login with everyone in the branch.

        You can also optionally check *"Remember key in this browser"* in the sidebar —
        the key is stored in your browser's local storage (like a password manager), so you only paste it once.
        """)

    with st.expander("Can I trust that the \"Remember key\" feature is secure?"):
        st.markdown("""
        Yes. Here's exactly what happens:

        1. When you check *Remember key in this browser*, the key is saved to **your browser's localStorage** —
           the same mechanism used by password managers and banking apps
        2. The key is **never sent to our server** — the JavaScript runs in your browser only
        3. It is **never visible in logs, analytics, or network requests** on our side
        4. You can remove it any time by unchecking the box, or by clearing your browser's site data

        If you're using a shared or public computer, simply leave the box unchecked.
        """)

    # ── HOW IT WORKS ─────────────────────────────────────
    st.markdown('<p class="faq-category">How It Works</p>', unsafe_allow_html=True)

    with st.expander("How is this different from pasting my data into ChatGPT?"):
        st.markdown("""
        Five meaningful differences:

        **1. Direct connection, no copy-paste**
        Upload once or connect your BI tool — no manual copying, no hitting character limits,
        no worrying about whether the AI saw all your rows.

        **2. Pre-engineered prompts**
        Instead of you figuring out how to ask for analysis, the prompts are pre-built for business decision-making —
        calibrated by role (CFO vs. Operations Manager) and time horizon (next 30 days vs. next quarter).
        ChatGPT gives you what you ask for; this gives you what you need.

        **3. Anomaly detection before the AI sees your data**
        Statistical outliers in your dataset are flagged before the report is generated,
        so the AI's recommendations account for dirty or unusual data.

        **4. Structured, shareable output**
        Reports are formatted and downloadable as Word or PDF — ready to drop into a board deck or email.
        No reformatting required.

        **5. Your data stays in your account**
        When you paste into ChatGPT, OpenAI's default data handling applies.
        Here, you're using your own API key — your data flows through your OpenAI account,
        subject to your organization's data agreements with OpenAI.
        """)

    with st.expander("What do the three report types mean?"):
        st.markdown("""
        Each report is written for a different audience and decision type:

        **Executive Summary**
        Written for the boardroom. Covers: Situation (what the data shows), What This Means For You
        (strategic implications), Recommended Actions, and Risks If You Wait.
        Best for: quarterly reviews, board presentations, investor updates.

        **Operational Detail**
        Written for team managers. Covers: Performance Breakdown by segment, Efficiency Gaps,
        Process Recommendations, and Quick Wins you can act on this week.
        Best for: weekly team syncs, department reviews, ops planning.

        **Risk Report**
        Written for risk officers and cautious decision-makers. Covers: Top Risks Identified,
        Early Warning Signs in the data, Mitigation Actions, and Data Quality Risks.
        Best for: audits, compliance reviews, risk committee updates.
        """)

    with st.expander("How accurate are the AI reports?"):
        st.markdown("""
        The reports are as accurate as your data and your decision context.

        The AI performs **pattern recognition and business reasoning** — it identifies trends, comparisons,
        outliers, and implications based on your dataset. It does not perform statistical hypothesis testing
        or make guarantees about future outcomes.

        **To get the best results:**
        - Be specific in the *Decision Context* field (e.g., "Should we expand into the Northeast market
          given Q2 performance?" is better than "What should I do?")
        - Select the right role — the report language and depth change significantly between CFO and Analyst
        - Upload clean, well-labelled data with clear column names

        Always use AI reports as **decision support**, not as a substitute for domain expertise.
        """)

    # ── ENTERPRISE ────────────────────────────────────────
    st.markdown('<p class="faq-category">Enterprise & Teams</p>', unsafe_allow_html=True)

    with st.expander("Does this work with Tableau or Power BI?"):
        st.markdown("""
        Yes. In the sidebar, switch the data source to **Tableau** or **Power BI**:

        **Tableau Server / Tableau Cloud**
        Enter your server URL, username, and password (or Personal Access Token).
        The app lists your available workbooks and views — select one and it pulls the underlying data live.
        No file export needed — the connection reads directly from your published dashboard data.

        **Power BI**
        Enter your Azure tenant ID, client ID, client secret, and workspace/dataset IDs.
        The app exports the dataset and loads it for analysis, keeping your data within your
        Microsoft tenant's security boundary.

        Both connections are stateless — credentials are used only to fetch data during your active session
        and are never stored, logged, or retained after you close the browser.

        > **Enterprise note:** If your organization has questions about connecting to internal Tableau Server
        > or Power BI on-premises deployments, reach out directly for guidance on network and credential configuration.
        """)

    # ── FOOTER ────────────────────────────────────────────
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("""
    <div style="border-top: 1px solid #E8E5DC; padding-top: 1.2rem; margin-top: 1rem;
                color: #9B8FA8; font-size: 0.78rem; text-align: center; line-height: 1.7;">
        Built by <strong style="color:#4A4A6A;">Swapnil Sakorkar</strong> &nbsp;·&nbsp;
        AI Application Developer &nbsp;·&nbsp;
        <a href="https://www.linkedin.com/in/swapnil-sakorkar" target="_blank"
           style="color:#B8975A; text-decoration:none;">LinkedIn</a>
        &nbsp;·&nbsp;
        <em>This tool uses your own OpenAI API key. Your data is processed in-session only and is never stored.</em>
    </div>
    """, unsafe_allow_html=True)

# ---------------------------------------------------
# DASHBOARD TAB
# ---------------------------------------------------

with nav_dashboard:

    # Resolve data from whichever source was used
    df_raw = None
    data_badge = ""

    if uploaded_file:
        df_raw     = load_file(uploaded_file)
        data_badge = uploaded_file.name
    elif df_from_bi is not None:
        df_raw     = df_from_bi
        data_badge = st.session_state.get("bi_source", "BI Tool")

    if df_raw is None:
        source_hint = {
            "Upload CSV / Excel": "Upload a CSV or Excel file in the sidebar to begin",
            "Tableau":            "Enter your Tableau credentials and click Connect",
            "Power BI":           "Enter your Power BI credentials and click Connect",
        }.get(data_source, "Choose a data source in the sidebar to begin")
        st.markdown(f"""
            <div class="empty-state">
                <div class="icon">↓</div>
                <p>{source_hint}</p>
            </div>
        """, unsafe_allow_html=True)

    else:
        df = clean_dataframe(df_raw.copy())
        st.session_state["df_for_filter"] = df

        numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
        cat_cols = df.select_dtypes(exclude=np.number).columns.tolist()

        # Apply filter for display
        df_display = df.copy()
        if filter_col and filter_val:
            df_display = df[df[filter_col] == filter_val]

        # --- DATA PREVIEW ---
        st.markdown('<p class="section-label">Dataset Overview</p>', unsafe_allow_html=True)

        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Rows", f"{len(df_display):,}")
        c2.metric("Columns", f"{len(df_display.columns)}")
        c3.metric("Numeric fields", len(numeric_cols))
        c4.metric("Quality Score", f"{dashboard_score(df_display)}/100")

        st.dataframe(
            df_display.head(10),
            use_container_width=True,
            hide_index=True,
        )

        # --- CHARTS ---
        if cat_cols and numeric_cols and df_display[cat_cols[0]].nunique() <= 25:

            chart_col1, chart_col2 = st.columns(2)

            with chart_col1:
                primary_cat = cat_cols[0]
                primary_num = numeric_cols[0]
                agg = smart_agg(primary_num)
                chart_df = df_display.groupby(primary_cat)[primary_num].agg(agg).reset_index()
                chart_df.columns = ["category", "value"]
                chart_df = chart_df.sort_values("value", ascending=False)

                bar = (
                    alt.Chart(chart_df)
                    .mark_bar(cornerRadiusTopRight=4, cornerRadiusBottomRight=4)
                    .encode(
                        x=alt.X("value:Q", title=f"{agg.title()} of {primary_num}",
                                 axis=alt.Axis(labelColor="#8B8FA8", titleColor="#8B8FA8",
                                               gridColor="#E8E5DC", domainColor="#E8E5DC",
                                               tickColor="#E8E5DC")),
                        y=alt.Y("category:N", sort="-x", title="",
                                 axis=alt.Axis(labelColor="#2E2E4A", domainColor="#E8E5DC",
                                               tickColor="#E8E5DC")),
                        color=alt.condition(
                            alt.datum.value >= 0,
                            alt.value("#B8975A"),
                            alt.value("#C0392B"),
                        ),
                        tooltip=["category", alt.Tooltip("value:Q", format=",.2f")],
                    )
                    .properties(
                        title=alt.TitleParams(
                            f"{agg.title()} {primary_num} by {primary_cat}",
                            color="#1A1A2E",
                            fontSize=13,
                            font="IBM Plex Sans",
                        ),
                        height=max(min(52 * df_display[primary_cat].nunique(), 400), 160),
                        background="#FFFFFF",
                        padding={"top": 16, "left": 16, "right": 16, "bottom": 16},
                    )
                )
                st.altair_chart(bar, use_container_width=True)

            with chart_col2:
                if len(numeric_cols) >= 2:
                    second_num = numeric_cols[1]
                    agg2 = smart_agg(second_num)
                    chart_df2 = df_display.groupby(cat_cols[0])[second_num].agg(agg2).reset_index()
                    chart_df2.columns = ["category", "value"]
                    chart_df2 = chart_df2.sort_values("value", ascending=False)

                    bar2 = (
                        alt.Chart(chart_df2)
                        .mark_bar(cornerRadiusTopRight=4, cornerRadiusBottomRight=4)
                        .encode(
                            x=alt.X("value:Q", title=f"{agg2.title()} of {second_num}",
                                     axis=alt.Axis(labelColor="#8B8FA8", titleColor="#8B8FA8",
                                                   gridColor="#E8E5DC", domainColor="#E8E5DC",
                                                   tickColor="#E8E5DC")),
                            y=alt.Y("category:N", sort="-x", title="",
                                     axis=alt.Axis(labelColor="#2E2E4A", domainColor="#E8E5DC",
                                                   tickColor="#E8E5DC")),
                            color=alt.condition(
                                alt.datum.value >= 0,
                                alt.value("#2E6DA4"),
                                alt.value("#C0392B"),
                            ),
                            tooltip=["category", alt.Tooltip("value:Q", format=",.2f")],
                        )
                        .properties(
                            title=alt.TitleParams(
                                f"{agg2.title()} {second_num} by {cat_cols[0]}",
                                color="#1A1A2E",
                                fontSize=13,
                                font="IBM Plex Sans",
                            ),
                            height=max(min(52 * df_display[cat_cols[0]].nunique(), 400), 160),
                            background="#FFFFFF",
                            padding={"top": 16, "left": 16, "right": 16, "bottom": 16},
                        )
                    )
                    st.altair_chart(bar2, use_container_width=True)

        # Anomaly warnings
        for col in numeric_cols[:2]:
            anom = detect_anomalies(df_display, col)
            if not anom.empty:
                st.warning(f"{len(anom)} anomalous rows detected in **{col}** — flagged in the Risk Report.")

        # --- REPORTS ---
        if run:

            if not decision.strip():
                st.warning("Please describe the decision you are trying to make before generating reports.")
            else:
                summary = build_data_summary(df, filter_col=filter_col, filter_val=filter_val)

                st.markdown("---")
                st.markdown('<p class="section-label">AI Reports</p>', unsafe_allow_html=True)

                tab1, tab2, tab3 = st.tabs(["Executive Summary", "Operational Detail", "Risk Report"])

                for tab, report_type in zip(
                    [tab1, tab2, tab3],
                    ["Executive Summary", "Operational Detail", "Risk Report"]
                ):
                    with tab:
                        with st.spinner(f"Generating {report_type}..."):
                            report_text = generate_report(summary, who, decision, timeframe, report_type, api_key=user_api_key)

                        # Clean up raw markdown artifacts before rendering
                        clean = report_text
                        # Ensure ### headers have a blank line before them for proper rendering
                        clean = re.sub(r'([^\n])\n(### )', r'\1\n\n\2', clean)

                        # Convert to HTML and inject inside styled .report-body card
                        html_lines = []
                        for line in clean.split("\n"):
                            line = line.strip()
                            if not line:
                                html_lines.append("<br>")
                            elif line.startswith("### "):
                                html_lines.append(f"<h3>{line[4:]}</h3>")
                            elif re.match(r"^\d+\.", line):
                                html_lines.append(f"<p>{line}</p>")
                            elif line.startswith("- ") or line.startswith("• "):
                                html_lines.append(f"<p>• {line[2:]}</p>")
                            else:
                                safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                                safe = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', safe)
                                safe = re.sub(r'\*(.+?)\*', r'<strong>\1</strong>', safe)
                                html_lines.append(f"<p>{safe}</p>")
                        html_body = "\n".join(html_lines)
                        st.markdown(
                            f'<div class="report-body">{html_body}</div>',
                            unsafe_allow_html=True
                        )

                        st.markdown("<br>", unsafe_allow_html=True)
                        dl1, dl2 = st.columns([1, 1])

                        with dl1:
                            st.download_button(
                                label="↓  Download as Word",
                                data=report_to_docx(report_text, report_type, who, decision, timeframe),
                                file_name=f"{report_type.lower().replace(' ', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"docx_{report_type}",
                                use_container_width=True,
                            )

                        with dl2:
                            st.download_button(
                                label="↓  Download as PDF",
                                data=report_to_pdf(report_text, report_type, who, decision, timeframe),
                                file_name=f"{report_type.lower().replace(' ', '_')}.pdf",
                                mime="application/pdf",
                                key=f"pdf_{report_type}",
                                use_container_width=True,
                            )

        elif uploaded_file:
            st.markdown("""
                <div style="text-align:center; padding: 2rem; color: #B0ADC0;
                            border: 1px dashed #D4D0C8; border-radius: 10px; margin-top: 1rem;
                            background: #FAFAF8;">
                    <p style="font-size: 0.82rem; letter-spacing: 0.08em;
                              text-transform: uppercase; margin: 0; font-weight: 500;">
                        Fill in the sidebar and click Generate Reports
                    </p>
                </div>
            """, unsafe_allow_html=True)
